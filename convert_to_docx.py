from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from indigo import Indigo
indigo = Indigo()
from indigo.renderer import IndigoRenderer
ir = IndigoRenderer(indigo)
indigo.setOption("render-output-format", "png")

#модуль для сохранения списка reactions в виде вордовского документа
work_dir = '/media/oleg/second_ssd/temp_files_parser'

def draw_rxn_add_to_doc(smiles, document, work_dir = work_dir):
    rxn = indigo.loadReaction(smiles) 
    f_name = os.path.join(work_dir, 'temp_fig.png')
    ir.renderToFile(rxn, f_name)
    document.add_picture(f_name, width=Inches(6))
    return True

def draw_mol_add_to_doc(smiles, document, work_dir = work_dir):
    mol = indigo.loadMolecule(smiles)
    f_name = os.path.join(work_dir, 'temp_fig_mol.png')
    ir.renderToFile(mol, f_name)
    document.add_picture(f_name)
    return True

def get_doc_reactions(reactions, doi, document):
    
    f_mol_name = os.path.join(work_dir, 'temp_fig_mol.png')
    
#     document = Document()
    
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Reactions report for ' + doi)
    run.font.size = Pt(14)
    run.bold = True
    
    for idx, reaction in enumerate(reactions):
        try:
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Reaction No. ' + str(idx))
            run.font.size = Pt(13)
            run.bold = True
            draw_rxn_add_to_doc(reaction['smiles'], document) #нарисовали реакцию

            paragraph = document.add_paragraph()
            run = paragraph.add_run('Protocol:')
            run.font.size = Pt(12)
            run.italic = True

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(reaction['procedure'])
            run.font.size = Pt(10)

            paragraph = document.add_paragraph()
            run = paragraph.add_run('Initial Protocol:')
            run.font.size = Pt(12)
            run.italic = True

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(reaction['initial_protocol'])
            run.font.size = Pt(10)

            paragraph = document.add_paragraph()
            run = paragraph.add_run('Conditions:')
            run.font.size = Pt(12)
            run.italic = True

            if len(reaction['conditions']) > 0:
                for cond in reaction['conditions']:
                    if cond is not None:
                        p = document.add_paragraph()
                        fmt = p.paragraph_format
                        fmt.line_spacing = 1
                        fmt.space_after = 1
                        fmt.space_before = 1
                        for k,v in cond.items():
                            run = p.add_run(str(k) + ': ' + str(v))
                            run.font.size = Pt(10)
            else:
                p = document.add_paragraph()
                fmt = p.paragraph_format
                fmt.line_spacing = 1
                fmt.space_after = 1
                fmt.space_before = 1
                run = p.add_run('None')
                run.font.size = Pt(10)

            #scale
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Scale: ' + str(reaction['scale']))
            run.font.size = Pt(12)

            #compounds
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Compounds:')
            run.font.size = Pt(12)
            run.italic = True

            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Molecule'
            hdr_cells[1].text = 'Compound_role'
            hdr_cells[2].text = 'Amounts'
            hdr_cells[3].text = 'Score'

            for comp in reaction['compounds']:
                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()

                mol = indigo.loadMolecule(comp['smiles'])
                ir.renderToFile(mol, f_mol_name)
                run.add_picture(f_mol_name, width =Inches(min([1.5, mol.countAtoms()/8])))

                row_cells[1].text = comp['compound_role']
                row_cells[2].text = str(comp['digit_amounts'])
                row_cells[3].text = str(comp['score'])

            #scores
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Scores:')
            run.font.size = Pt(12)
            run.italic = True

            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'score_name'
            hdr_cells[1].text = 'value'
            for k,v in reaction['score_dict'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = k
                row_cells[1].text = str(v)
        except Exception as e:
            print(e)
            pass
#     document.save(file_path)